from __future__ import annotations

import csv
import io
import re
from pathlib import Path
from typing import Iterable

import openpyxl
import pdfplumber
import xlrd
from docx import Document

from .models import Job, now_iso
from .text import clean, stable_id


HEADER_ALIASES = {
    "organization": ("招聘单位", "招录机关", "单位名称", "用人单位", "职位所在单位"),
    "code": ("岗位代码", "职位代码", "职位编号", "岗位编号"),
    "position": ("岗位名称", "职位名称", "岗位", "职位简介"),
    "city": ("地区名称", "工作地点", "职位所在地", "所在地"),
    "education": ("学历", "学历要求"),
    "degree": ("学位", "学位要求"),
    "majors": ("专业", "专业要求", "所学专业"),
    "political_status": ("政治面貌",),
    "target_group": ("招聘对象", "招录对象", "人员性质", "对象"),
    "experience": ("基层工作经历", "工作经历", "相关工作经历"),
    "other_requirements": ("其他条件", "其他条件和说明", "其它", "备注", "有关要求", "资格条件"),
}


def _map_headers(headers: list[str]) -> dict[str, int]:
    mapped: dict[str, int] = {}
    used_columns: set[int] = set()
    for index, header in enumerate(headers):
        normalized = clean(header).replace(" ", "")
        for field, aliases in HEADER_ALIASES.items():
            if field not in mapped and index not in used_columns and any(alias == normalized for alias in aliases):
                mapped[field] = index
                used_columns.add(index)
                break
    for index, header in enumerate(headers):
        normalized = clean(header).replace(" ", "")
        for field, aliases in HEADER_ALIASES.items():
            if field not in mapped and index not in used_columns and any(alias in normalized for alias in aliases):
                if field == "position" and "代码" in normalized:
                    continue
                mapped[field] = index
                used_columns.add(index)
                break
    return mapped


def _find_header(rows: list[list[object]]) -> tuple[int, dict[str, int]]:
    best = (0, {})
    for index, row in enumerate(rows[:20]):
        mapping = _map_headers([clean(v) for v in row])
        if len(mapping) > len(best[1]):
            best = (index, mapping)
    single_best_size = len(best[1])
    for index, row in enumerate(rows[:19]):
        if index + 1 < min(len(rows), 20):
            width = max(len(row), len(rows[index + 1]))
            upper = [clean(row[i]) if i < len(row) else "" for i in range(width)]
            carried = ""
            for i, value in enumerate(upper):
                if value:
                    carried = value
                else:
                    upper[i] = carried
            lower = [clean(rows[index + 1][i]) if i < len(rows[index + 1]) else "" for i in range(width)]
            combined = [clean(f"{upper[i]} {lower[i]}") for i in range(width)]
            combined_mapping = _map_headers(combined)
            if len(combined_mapping) > max(len(best[1]), single_best_size):
                best = (index + 1, combined_mapping)
    return best


def _sheet_rows(path: Path) -> Iterable[tuple[str, list[list[object]]]]:
    suffix = path.suffix.lower()
    if suffix == ".xls":
        book = xlrd.open_workbook(path)
        try:
            for sheet in book.sheets():
                yield sheet.name, [sheet.row_values(i) for i in range(sheet.nrows)]
        finally:
            book.release_resources()
    elif suffix == ".xlsx":
        book = openpyxl.load_workbook(path, read_only=True, data_only=True)
        try:
            for sheet in book.worksheets:
                yield sheet.title, [list(row) for row in sheet.iter_rows(values_only=True)]
        finally:
            book.close()
    elif suffix == ".csv":
        text = path.read_text(encoding="utf-8-sig", errors="replace")
        yield path.stem, list(csv.reader(io.StringIO(text)))


def jobs_from_spreadsheet(path: Path, meta: dict) -> list[Job]:
    jobs: list[Job] = []
    for sheet_name, rows in _sheet_rows(path):
        if not rows:
            continue
        header_index, mapping = _find_header(rows)
        if len(mapping) < 3 or "position" not in mapping:
            continue
        previous: dict[str, str] = {}
        for row_index, raw_row in enumerate(rows[header_index + 1 :], start=header_index + 2):
            row = [clean(v) for v in raw_row]
            values: dict[str, str] = {}
            for field, column in mapping.items():
                value = row[column] if column < len(row) else ""
                if not value and field in {"organization", "city"}:
                    value = previous.get(field, "")
                values[field] = value
                if value:
                    previous[field] = value
            position = values.get("position", "")
            organization = values.get("organization", "")
            if not position or position in {"岗位名称", "职位名称", "职位"}:
                continue
            code = values.get("code", "")
            title = f"{organization} - {position}".strip(" -")
            raw_city = values.get("city", "")
            known_city = next((city for city in ("苏州", "南京", "无锡", "南通", "常州", "扬州", "镇江", "徐州", "盐城", "泰州", "淮安", "宿迁", "连云港") if city in raw_city), "")
            category = meta.get("category", "身份待核实")
            if category == "公务员" and "参照管理" in f"{position} {values.get('other_requirements', '')}":
                category = "参公"
            job = Job(
                id=stable_id(meta.get("source_url", str(path)), sheet_name, code, title),
                title=title,
                organization=organization,
                position=position,
                category=category,
                employment_status=meta.get("employment_status", "待核实"),
                city=known_city or meta.get("city", "江苏"),
                source_name=meta.get("source_name", path.name),
                source_url=meta.get("source_url", ""),
                source_file=path.name,
                source_sheet=sheet_name,
                source_row=row_index,
                source_code=code,
                official=meta.get("official", True),
                published_at=meta.get("published_at", ""),
                deadline=meta.get("deadline", ""),
                education=values.get("education", ""),
                degree=values.get("degree", ""),
                majors=values.get("majors", ""),
                political_status=values.get("political_status", ""),
                target_group=values.get("target_group", ""),
                experience=values.get("experience", ""),
                other_requirements=values.get("other_requirements", ""),
                attachment_urls=meta.get("attachment_urls", []),
                summary=(
                    f"原始职位表：{path.name}；工作表：{sheet_name}；第{row_index}行"
                    + (f"；代码：{code}" if code else "")
                ),
                discovered_at=now_iso(),
                last_seen_at=now_iso(),
            )
            jobs.append(job)
    return jobs


def text_from_attachment(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        with pdfplumber.open(path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    if suffix == ".docx":
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            paragraphs.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(paragraphs)
    return ""
